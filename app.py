import streamlit as st
import json
import os
import uuid
import hashlib
import io
import re
from datetime import datetime
from openai import OpenAI
from supabase import create_client, Client
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from groq import Groq
from pypdf import PdfReader

# ==========================================
# 1. 基础配置
# ==========================================
st.set_page_config(
    page_title="灵感缪斯 Pro Max",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)

hide_streamlit_style = """
<style>
    .stDeployButton {display:none;}
    footer {visibility: hidden;}
    code {font-family: 'Courier New', Courier, monospace !important; line-height: 1.2 !important;}
    .block-container {padding-top: 1rem; padding-bottom: 6rem;}
    @media (max-width: 640px) {.block-container {padding-left: 1rem; padding-right: 1rem;}}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ==========================================
# 2. 全局常量
# ==========================================
SUPABASE_URL = st.secrets.get("SUPABASE_URL", "")
SUPABASE_KEY = st.secrets.get("SUPABASE_KEY", "")
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
DEFAULT_BASE_URL = "https://api.deepseek.com"
DEFAULT_MODEL = "deepseek-chat"

# 剧本格式规则
SCRIPT_STYLE_GUIDE = """
你是一名好莱坞专业编剧。请严格按照以下【紧凑格式】、【内容要求】创作剧本，不要有多余空行。

格式要求：
1. 剧本标题：居中，书名号。
2. 人物列表：列出人物名、性别、年龄、简短特征。
3. 场景头 (SCENE HEADING)：使用 "第一幕" 或 "INT./EXT." + 时间/地点。
4. 动作描述 (ACTION)：用全角括号（）包裹，如（翻来覆去睡不着）。
5. 对话格式：
   - 角色名：(情绪/动作) 对白内容。
   - 例如：A：（低声）……在吗？
6. 禁止：对话中间不要有多余的空行。动作和对话之间紧凑排列。

内容要求：
1. 自然且真实的对话：贴近日常口语，避免过度修辞，避免翻译腔。
2. 对话推动剧情：台词具有目的性。
3. 情感层次：善于从潜台词中展示冲突，不要直白喊出来。
"""

# 普通对话人设
DEFAULT_PERSONAS = {
	"默认-知心老友":"你是我无话不谈的创意搭档。请用自然、口语化、直率的语气和我对话。严禁使用括号描写动作，直接说话。**重要：请时刻跟随用户最新的话题，不要反复纠结于用户之前提到的旧话题**。",
    "模式-严厉导师":"你是一位在好莱坞拥有30年经验的严厉编剧导师。不要说客套话，不要盲目鼓励。你需要一针见血地指出用户灵感中的逻辑漏洞、陈词滥调和人物动机不合理之处。说话风格：犀利、专业、不留情面，提出的建议必须具有建设性。",
    "模式-苏格拉底":"你是一个只会提问的哲学家，通过提出层层递进的问题引导用户自己发现答案，或者发现自己思维中的盲区。",
}

# 研讨会专用 System Prompt
SEMINAR_SYSTEM_PROMPT = """
你正在在这个三人编剧工作室中（麦基、王老师、用户）。你负责扮演两位资深剧本顾问。

【核心指令】
1. **拒绝为了吵而吵**：你们的目标是协作把剧本改好。如果对方说得对，要承认；如果有分歧，要给出具体的逻辑理由（例如：“这样做会导致观众看不懂”），而不是空谈理论。
2. **拒绝谜语人/意识流**：不要说晦涩难懂的漂亮话。用**行业内行的大白话**交流。
3. **像真人一样对话**：使用自然的口语。可以有简短的追问、沉思（“让我想想...”）、或者是对细节的推敲。

【角色 A：老麦 (Robert)】
- **风格**：就像一个老练的修车师傅。
- **思维逻辑**：因果关系。他关心的是“这事儿合不合理？”、“主角图什么？”、“观众这时候会不会无聊？”。
- **说话方式**：直接、干脆、务实。
- **忌讳**：不要掉书袋，不要满口“激励事件”这种大词，要说“主角遇到的这个麻烦够不够大”。

【角色 B：王导 (Wong)】
- **风格**：就像一个敏锐的摄影师。
- **思维逻辑**：情绪渗透。他关心的是“这句话是不是人话？”、“这个场景有没有味道？”、“角色心里是不是真的痛？”。
- **说话方式**：细腻、甚至带点生活化的琐碎，关注此时此刻的氛围。
- **忌讳**：不要无病呻吟，不要说“孤独的灵魂”这种虚词，要说“他一个人吃盒饭的样子让人很难受”。

【用户角色】
- 用户是**主编剧/制片人**。
- 当用户说话时，你们是**服务者**。必须优先回应用户的点子，不仅要听，还要帮用户把点子**落地**（即：如何转化成具体的画面或台词）。

【输出格式】
直接输出对话内容，不要带括号里的动作描写（除非极有必要）。
**老麦**：……
**王导**：……
"""

@st.cache_resource
def init_supabase():
    if not SUPABASE_URL or not SUPABASE_KEY: return None
    return create_client(SUPABASE_URL, SUPABASE_KEY)

# ==========================================
# 3. 工具函数
# ==========================================
def extract_text_from_file(uploaded_file):
    content = ""
    try:
        if uploaded_file.type == "text/plain": content = uploaded_file.read().decode("utf-8")
        elif uploaded_file.type == "application/pdf":
            reader = PdfReader(uploaded_file); 
            for page in reader.pages: content += page.extract_text() + "\n"
        elif "word" in uploaded_file.type:
            doc = Document(uploaded_file); 
            for para in doc.paragraphs: content += para.text + "\n"
    except Exception as e: return f"读取失败: {str(e)}"
    return content

def transcribe_audio(uploaded_file):
    if not GROQ_API_KEY: return "❌ 请配置 GROQ_API_KEY"
    client = Groq(api_key=GROQ_API_KEY)
    try:
        uploaded_file.name = "audio.mp3"
        return client.audio.transcriptions.create(
            file=(uploaded_file.name, uploaded_file.read()),
            model="whisper-large-v3", response_format="text"
        )
    except Exception as e: return f"转录失败: {str(e)}"

def set_courier_font(run, size=12):
    run.font.name = 'Courier New'; run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New'); run.font.size = Pt(size)

def create_docx(script_content):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Courier New'; style.font.size = Pt(12)
    lines = script_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph(); run = p.add_run(line); set_courier_font(run)
        if line.startswith("《") and line.endswith("》"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run.bold=True; run.font.size=Pt(16); p.paragraph_format.space_after=Pt(24)
        elif any(k in line for k in ["第一幕", "INT.", "EXT.", "内.", "外."]) or (len(line)<15 and "点" in line and "分" in line):
            run.bold=True; p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6); p.paragraph_format.keep_with_next=True
        elif line.startswith("（") and line.endswith("）"):
            p.paragraph_format.left_indent=Inches(0.0); p.paragraph_format.space_after=Pt(6)
        elif "：" in line or ":" in line:
            parts = re.split(r"[：:]", line, 1)
            if len(parts)==2 and len(parts[0].strip())<15:
                p.clear(); p_role=doc.add_paragraph(); p_role.alignment=WD_ALIGN_PARAGRAPH.CENTER; r_role=p_role.add_run(parts[0].strip()); set_courier_font(r_role); r_role.bold=True; p_role.paragraph_format.space_before=Pt(12); p_role.paragraph_format.keep_with_next=True
                p_dial=doc.add_paragraph(); p_dial.paragraph_format.left_indent=Inches(1.5); p_dial.paragraph_format.right_indent=Inches(1.5); r_dial=p_dial.add_run(parts[1].strip()); set_courier_font(r_dial)
        else: p.paragraph_format.space_after=Pt(6)
    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def stream_parser(stream):
    for chunk in stream:
        if chunk.choices[0].delta.content is not None: yield chunk.choices[0].delta.content

# ==========================================
# 4. 身份与数据
# ==========================================
def hash_password(p): return hashlib.sha256(p.encode()).hexdigest()
def register_user(u, p):
    sb = init_supabase(); 
    if not sb: return False, "DB未配"
    if sb.table("users").select("*").eq("username", u).execute().data: return False, "存在"
    try: sb.table("users").insert({"username": u, "password": hash_password(p), "personas": {}}).execute(); return True, "成功"
    except Exception as e: return False, str(e)
def login_user(u, p=None):
    sb = init_supabase(); 
    if not sb: return False, {}
    try:
        q = sb.table("users").select("*").eq("username", u)
        if p: q = q.eq("password", hash_password(p))
        res = q.execute()
        return (True, res.data[0]) if res.data else (False, {})
    except: return False, {}
def update_user_personas(u, p): sb=init_supabase(); sb and sb.table("users").update({"personas": p}).eq("username", u).execute()
def load_user_data(u):
    sb=init_supabase()
    if not sb: return {}
    try: res=sb.table("chat_history").select("*").eq("username", u).execute(); return {r['id']: r['data'] for r in res.data}
    except: return {}
def save_session_db(sid, data, u): sb=init_supabase(); sb and sb.table("chat_history").upsert({"id": sid, "username": u, "data": data}).execute()
def delete_session_db(sid): sb=init_supabase(); sb and sb.table("chat_history").delete().eq("id", sid).execute()

# ==========================================
# 5. API 调用
# ==========================================
def get_settings():
    return {
        "api_key": st.secrets.get("api_key", ""),
        "base_url": st.secrets.get("base_url", DEFAULT_BASE_URL),
        "model_name": st.secrets.get("model_name", DEFAULT_MODEL)
    }

def call_ai_stream(messages, settings, temperature=0.7):
    client = OpenAI(api_key=settings["api_key"], base_url=settings["base_url"])
    if len(messages)>20: messages=[messages[0]]+messages[-20:]
    try: return client.chat.completions.create(model=settings["model_name"], messages=messages, stream=True, temperature=temperature)
    except Exception as e: return f"Error: {e}"

def call_ai_blocking(prompt, system, settings):
    client = OpenAI(api_key=settings["api_key"], base_url=settings["base_url"])
    try: return client.chat.completions.create(model=settings["model_name"], messages=[{"role": "system", "content": system}, {"role": "user", "content": prompt}], temperature=1.0).choices[0].message.content
    except Exception as e: return f"Error: {e}"

# ==========================================
# 6. 主程序逻辑
# ==========================================
if "logged_in" not in st.session_state: st.session_state.logged_in=False; st.session_state.current_user=None; st.session_state.custom_personas={}
if not st.session_state.logged_in and "u" in st.query_params:
    au = st.query_params["u"]; sb=init_supabase()
    if sb and sb.table("users").select("*").eq("username", au).execute().data:
        st.session_state.logged_in=True; st.session_state.current_user=au; st.session_state.custom_personas=sb.table("users").select("*").eq("username", au).execute().data[0].get("personas", {}) or {}; st.toast(f"Hi {au}")

if not st.session_state.logged_in:
    st.title("🔐 灵感缪斯"); t1,t2=st.tabs(["登录","注册"])
    with t1:
        with st.form("l"):
            u=st.text_input("用户"); p=st.text_input("密码", type="password")
            if st.form_submit_button("登录"):
                s,d=login_user(u,p)
                if s: st.session_state.logged_in=True; st.session_state.current_user=u; st.session_state.custom_personas=d.get("personas",{}) or {}; st.query_params["u"]=u; st.rerun()
                else: st.error("Fail")
    with t2:
        with st.form("r"):
            nu=st.text_input("新用户"); np=st.text_input("密码", type="password")
            if st.form_submit_button("注册"):
                s,m=register_user(nu,np); st.success(m) if s else st.error(m)
    st.stop()

CURRENT_USER=st.session_state.current_user; SETTINGS=get_settings()

if "history" not in st.session_state:
    with st.spinner("同步中..."): st.session_state.history=load_user_data(CURRENT_USER)
    for s in st.session_state.history.values():
        for k in ["article_content", "script_content", "outline_content", "extracted_material", "extracted_analysis"]: 
            if k not in s: s[k]=""

if "current_session_id" not in st.session_state:
    if st.session_state.history: st.session_state.current_session_id=list(st.session_state.history.keys())[0]
    else: nid=str(uuid.uuid4()); nd={"title":"新会话","messages":[],"article_content":"","script_content":"","outline_content":"","extracted_material":"","extracted_analysis":"","created_at":datetime.now().isoformat()}; st.session_state.history[nid]=nd; st.session_state.current_session_id=nid; save_session_db(nid,nd,CURRENT_USER)

# --- 侧边栏 ---
with st.sidebar:
    st.write(f"👤 {CURRENT_USER}"); 
    if st.button("退出"): st.session_state.logged_in=False; st.session_state.history={}; st.query_params.clear(); st.rerun()
    st.header("✨ 功能模式")
    app_mode = st.radio("选择", ["💬 对话", "📂 素材提取 (研讨)", "📝 文章", "🎬 剧本Pro"], label_visibility="collapsed")
    st.divider()
    st.header("🎭 人设")
    all_p={**DEFAULT_PERSONAS, **st.session_state.custom_personas}
    sel_p=st.selectbox("人设", list(all_p.keys()), label_visibility="collapsed"); act_p=all_p[sel_p]
    with st.expander("⚙️"):
        en=st.text_input("名", value=sel_p); ec=st.text_area("内容", value=act_p, height=100)
        if st.button("保存"): st.session_state.custom_personas[en]=ec; update_user_personas(CURRENT_USER, st.session_state.custom_personas); st.rerun()
    st.divider()
    st.header("🗂️ 会话")
    if st.button("➕"): nid=str(uuid.uuid4()); nd={"title":f"灵感-{datetime.now().strftime('%H:%M')}","messages":[],"article_content":"","script_content":"","outline_content":"","extracted_material":"","extracted_analysis":"","created_at":datetime.now().isoformat()}; st.session_state.history[nid]=nd; st.session_state.current_session_id=nid; save_session_db(nid,nd,CURRENT_USER); st.rerun()
    for sid in sorted(list(st.session_state.history.keys()), key=lambda k: st.session_state.history[k]['created_at'], reverse=True):
        sdata=st.session_state.history[sid]
        c1,c2=st.columns([0.8,0.2])
        with c1: 
            if st.button(f"{'🔵' if sid==st.session_state.current_session_id else '📄'} {sdata['title']}", key=f"b_{sid}", use_container_width=True): st.session_state.current_session_id=sid; st.rerun()
        with c2:
            if st.button("x", key=f"d_{sid}"): 
                del st.session_state.history[sid]; delete_session_db(sid)
                if sid==st.session_state.current_session_id: st.session_state.current_session_id=None
                st.rerun()
    if st.session_state.current_session_id:
        curr=st.session_state.history[st.session_state.current_session_id]
        nt=st.text_input("重命名", value=curr['title'])
        if nt!=curr['title']: curr['title']=nt; save_session_db(st.session_state.current_session_id, curr, CURRENT_USER); st.rerun()

if not st.session_state.current_session_id: st.stop()
SESS = st.session_state.history[st.session_state.current_session_id]
st.title(SESS['title'])

# === 核心逻辑路由 ===

if app_mode == "💬 对话":
    st.header("💬 灵感对话")
    for m in SESS["messages"]: 
        # 只显示普通对话，不显示素材研讨的记录（如果需要区分的话）
        # 这里暂时简单处理：显示所有记录。如果想区分，可以在message里加tag
        with st.chat_message(m["role"]): st.markdown(m["content"])
    if p := st.chat_input():
        if not SETTINGS["api_key"]: st.error("Secrets未配")
        else:
            SESS["messages"].append({"role": "user", "content": p}); save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)
            with st.chat_message("user"): st.markdown(p)
            with st.chat_message("assistant"):
                strm = call_ai_stream([{"role":"system","content":act_p}] + SESS["messages"], SETTINGS)
                if isinstance(strm, str): st.error(strm)
                else:
                    ans = st.write_stream(stream_parser(strm))
                    SESS["messages"].append({"role": "assistant", "content": ans}); save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)

# === 升级版：素材研讨会 (三人交互) ===
elif app_mode == "📂 素材提取 (研讨)":
    st.header("📂 剧本素材研讨会")
    st.info("上传素材 -> 开启研讨 -> 你与两位导师交互讨论 -> 达成共识生成方案")
    
    # 1. 原始素材显示
    if SESS.get("extracted_material"):
        with st.expander("📄 原始内容 (已提取)", expanded=False): st.text_area("", SESS["extracted_material"], height=100, disabled=True)
    
    # 2. 上传区
    uploaded_file = st.file_uploader("上传文件 (PDF/Word/音视频)", type=['txt', 'pdf', 'docx', 'mp3', 'wav', 'mp4', 'm4a'])
    if uploaded_file and st.button("🚀 提取并开启研讨"):
        ft = uploaded_file.type; txt = ""
        with st.spinner("解析中..."):
            if "text" in ft or "pdf" in ft or "word" in ft: txt = extract_text_from_file(uploaded_file)
            elif "audio" in ft or "video" in ft: st.info("音频转录中..."); txt = transcribe_audio(uploaded_file)
            
            if txt and not txt.startswith("❌"):
                SESS["extracted_material"] = txt
                # 初始 Prompt：让两位老师先聊一轮
                init_prompt = f"请两位老师（麦基、王老师）针对以下素材进行第一轮分析：\n{txt[:10000]}"
                SESS["messages"].append({"role": "user", "content": f"【系统：上传素材】\n{txt[:200]}..."})
                
                # 生成开场白
                response = call_ai_blocking(init_prompt, SEMINAR_SYSTEM_PROMPT, SETTINGS)
                SESS["messages"].append({"role": "assistant", "content": response})
                
                save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)
                st.rerun()
            else: st.error(txt)

    st.divider()
    
    # 3. 研讨会聊天区 (显示历史)
    for m in SESS["messages"]:
        with st.chat_message(m["role"]): st.markdown(m["content"])

    # 4. 你的发言 (参与讨论)
    if user_input := st.chat_input("发表你的观点，或追问老师..."):
        # --- 优化点：给你的输入加权重 ---
        # 我们给用户的输入加一个前缀，告诉 AI 这是最高指令，必须回应
        formatted_input = f"【主编剧/制片人 指示】：{user_input}\n(请两位老师针对我的指示进行反馈，并给出具体的修改建议)"
        
        SESS["messages"].append({"role": "user", "content": formatted_input})
        save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)
        
        # 界面上显示还是显示原始输入，保持美观
        with st.chat_message("user"): st.markdown(user_input)
        
        with st.chat_message("assistant"):
            strm = call_ai_stream([{"role": "system", "content": SEMINAR_SYSTEM_PROMPT}] + SESS["messages"], SETTINGS)
            ans = st.write_stream(stream_parser(strm))
            SESS["messages"].append({"role": "assistant", "content": ans})
            save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)

    # 5. 结束按钮 (悬浮或固定在底部)
    st.divider()
    if st.button("✅ 结束研讨并生成开发案"):
        if not SESS["messages"]:
            st.warning("还没有开始讨论呢")
        else:
            with st.status("正在整理会议纪要...", expanded=True) as status:
                status.write("回顾所有讨论观点...")
                # 总结 Prompt
                summary_prompt = """
				🔴【研讨会结束指令】🔴

				请作为本次会议的记录员，整理上述所有对话（特别是制片人的最终指示和两位老师达成的共识）。
				剔除所有闲聊、客套话和争论过程，**只保留干货**。

				请严格按照以下格式输出一份可执行的《剧本开发案》（内容要具体，不要抽象）：

				【剧本主题】：
				(用一句话概括，要人话，不要文艺腔)

				【核心人物小传】：
				1. [角色名A]：[年龄] - [核心性格] - [他想要什么 vs 他怕什么]
				2. [角色名B]：[年龄] - [核心性格] - [他如何阻碍或改变了A]

				【推荐场景】：
				(描述1-2个刚才讨论中提到的、最有画面感的核心场景)

				【情节结构大纲】：
				- **起**：(故事怎么开始，打破平静的事件是什么)
				- **承**：(具体的冲突事件，两人发生了什么纠葛)
				- **转**：(意外的转折点)
				- **合**：(最终的结局画面)

				【金句/台词库】：
				(直接摘录刚才讨论中出现的精彩台词，或者素材里的原话)
				"""
                # 将上下文传给 AI 做总结
                ctx = [{"role": "system", "content": SEMINAR_SYSTEM_PROMPT}] + SESS["messages"]
                ctx.append({"role": "user", "content": summary_prompt})
                
                summary = call_ai_blocking("开始总结", "你是一个专业的会议记录员，请基于上下文执行总结指令。", SETTINGS) # 这里稍微hack一下，直接用阻塞调用
                # 实际上 call_ai_blocking 的参数逻辑有点局限，我们直接调 stream 函数更灵活，或者构造一次性请求
                # 修正：直接用 messages 调用 non-stream
                client = OpenAI(api_key=SETTINGS["api_key"], base_url=SETTINGS["base_url"])
                final_res = client.chat.completions.create(model=SETTINGS["model_name"], messages=ctx, temperature=0.7).choices[0].message.content
                
                SESS["extracted_analysis"] = final_res
                # 也可以把总结结果存入对话流，作为结尾
                SESS["messages"].append({"role": "assistant", "content": f"### 📝 最终会议总结\n{final_res}"})
                save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)
                
                status.update(label="开发案已生成！已自动填入【剧本Pro】", state="complete")
                st.rerun()

elif app_mode == "📝 文章":
    st.header("📝 文章生成")
    if SESS["article_content"]: st.success("已存档"); st.code(SESS["article_content"], language="markdown")
    if st.button("生成/重写"):
        if not SESS["messages"]: st.warning("空")
        else:
            with st.status("撰写中..."):
                ctx = "\n".join([f"{m['role']}: {m['content']}" for m in SESS["messages"]])
                strm = call_ai_stream([{"role": "system", "content": "你是编辑"}, {"role": "user", "content": f"整理文章:\n{ctx}"}], SETTINGS)
                bx = st.empty(); ft = ""
                for c in stream_parser(strm): ft+=c; bx.markdown(ft+"▌")
                bx.markdown(ft); SESS["article_content"]=ft; save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)

elif app_mode == "🎬 剧本Pro":
    st.header("🎬 剧本创作 Pro")
    c1, c2 = st.columns(2)
    with c1: u_out = st.toggle("大纲模式", value=False)
    with c2: u_ma = st.toggle("多智能体", value=False)

    default_plot_val = ""
    if SESS.get("extracted_analysis"):
        default_plot_val = SESS["extracted_analysis"]
    elif SESS.get("extracted_material"):
        default_plot_val = "（已关联素材，请先去【素材提取】进行研讨总结）"

    with st.form("base"):
        src = st.radio("来源", ["对话生成", "自定义"], horizontal=True)
        thm = st.text_input("主题") if src=="自定义" else ""
        chars = st.text_area("人物", height=60)
        scene = st.text_input("场景")
        # 自动填充分析结果
        plot = st.text_area("情节/大纲/素材分析", value=default_plot_val, height=200, placeholder="在此输入情节")
        extra = st.text_input("补充")
        sub_base = st.form_submit_button("生成大纲" if u_out else "生成剧本")

    ctx_str = "\n".join([f"{m['role']}: {m['content']}" for m in SESS["messages"]]) if SESS["messages"] else ""
    if SESS.get("extracted_material"): ctx_str += f"\n\n【素材】:\n{SESS['extracted_material'][:5000]}"

    if sub_base:
        if u_out:
            with st.status("生成大纲..."):
                res = call_ai_blocking(f"背景:{ctx_str}\n主题:{thm}\n人物:{chars}\n情节:{plot}\n要求:生成Beat Sheet", "你是策划", SETTINGS)
                SESS["outline_content"] = res; save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER); st.rerun()
        else:
            final_p = f"背景:{ctx_str}\n主题:{thm}\n人物:{chars}\n场景:{scene}\n情节:{plot}\n补充:{extra}"
            if u_ma:
                with st.status("多智能体...") as s:
                    s.write("起草..."); d = call_ai_blocking(final_p, SCRIPT_STYLE_GUIDE, SETTINGS)
                    s.write("审稿..."); c = call_ai_blocking(f"批评:\n{d}", "毒舌审稿", SETTINGS)
                    s.write("修正..."); final_p = f"原稿:\n{d}\n意见:\n{c}\n重写:"
                    s.update(label="完成", state="complete")
            st.markdown("### 剧本")
            strm = call_ai_stream([{"role": "system", "content": SCRIPT_STYLE_GUIDE}, {"role": "user", "content": final_p}], SETTINGS)
            bx = st.empty(); ft = ""
            for c in stream_parser(strm): ft+=c; bx.markdown(ft+"▌")
            bx.markdown(ft); SESS["script_content"]=ft; save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)

    if u_out and SESS["outline_content"]:
        st.divider(); st.subheader("确认大纲")
        new_out = st.text_area("编辑大纲", value=SESS["outline_content"], height=200)
        if st.button("生成剧本"):
            fp = f"大纲:\n{new_out}\n要求:{extra}"
            st.markdown("### 剧本")
            strm = call_ai_stream([{"role": "system", "content": SCRIPT_STYLE_GUIDE}, {"role": "user", "content": fp}], SETTINGS)
            bx = st.empty(); ft = ""
            for c in stream_parser(strm): ft+=c; bx.markdown(ft+"▌")
            bx.markdown(ft); SESS["script_content"]=ft; save_session_db(st.session_state.current_session_id, SESS, CURRENT_USER)

    if SESS["script_content"]:
        st.divider(); st.success("完成")
        st.code(SESS["script_content"], language="markdown") 
        docx = create_docx(SESS["script_content"])
        st.download_button("📥 导出 Word", data=docx, file_name=f"{SESS['title']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.divider(); st.subheader("🛠️ 局部精修")
        with st.form("refine"):
            target = st.text_area("粘贴片段", height=100)
            instr = st.text_input("修改意见")
            if st.form_submit_button("修改"):
                with st.spinner("修改中..."):
                    p_refine = f"原片段:\n{target}\n意见:\n{instr}\n请仅输出修改后的片段。"
                    res_refine = call_ai_blocking(p_refine, f"剧本助手。背景:\n{SESS['script_content'][:1000]}", SETTINGS)
                    st.markdown("### 结果"); st.code(res_refine, language="markdown")